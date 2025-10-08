"""
DOCX Export Service
Generates ATS-compatible Word documents from resume structure
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
import io

class DOCXExporter:
    """Generate ATS-compatible DOCX resumes"""

    def __init__(self):
        # ATS-safe font settings
        self.font_name = "Arial"
        self.font_size_body = Pt(11)
        self.font_size_name = Pt(18)
        self.font_size_heading = Pt(12)
        self.font_size_subheading = Pt(11)

    def generate_docx(self, resume_structure: Dict[str, Any]) -> bytes:
        """
        Generate DOCX file from resume structure

        Args:
            resume_structure: Resume data structure

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Set document margins (0.75 inches all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Add contact information
        self._add_contact_info(doc, resume_structure.get('contact_info', {}))

        # Add professional summary
        self._add_section(doc, "PROFESSIONAL SUMMARY")
        self._add_paragraph(doc, resume_structure.get('summary', ''))

        # Add professional experience
        self._add_section(doc, "PROFESSIONAL EXPERIENCE")
        for exp in resume_structure.get('experience', []):
            self._add_experience(doc, exp)

        # Add skills
        skills = resume_structure.get('skills', {})
        if skills:
            self._add_section(doc, "SKILLS")
            for category, skill_list in skills.items():
                self._add_skills_category(doc, category, skill_list)

        # Add education
        education = resume_structure.get('education', [])
        if education:
            self._add_section(doc, "EDUCATION")
            for edu in education:
                self._add_education(doc, edu)

        # Add certifications
        certifications = resume_structure.get('certifications', [])
        if certifications:
            self._add_section(doc, "CERTIFICATIONS")
            for cert in certifications:
                self._add_certification(doc, cert)

        # Convert to bytes
        docx_stream = io.BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)

        return docx_stream.getvalue()

    def _add_contact_info(self, doc: Document, contact: Dict) -> None:
        """Add contact information section"""
        # Name (centered, large)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(contact.get('name', ''))
        name_run.font.size = self.font_size_name
        name_run.font.name = self.font_name
        name_run.bold = True

        # Contact details (centered)
        contact_parts = []
        if contact.get('phone'):
            contact_parts.append(contact['phone'])
        if contact.get('email'):
            contact_parts.append(contact['email'])

        if contact_parts:
            contact_line1 = doc.add_paragraph()
            contact_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run1 = contact_line1.add_run(' | '.join(contact_parts))
            contact_run1.font.size = self.font_size_body
            contact_run1.font.name = self.font_name

        # Location and links
        contact_parts2 = []
        if contact.get('location'):
            contact_parts2.append(contact['location'])

        if contact_parts2:
            contact_line2 = doc.add_paragraph()
            contact_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run2 = contact_line2.add_run(' | '.join(contact_parts2))
            contact_run2.font.size = self.font_size_body
            contact_run2.font.name = self.font_name

        # LinkedIn and portfolio (separate line)
        links = []
        if contact.get('linkedin'):
            links.append(contact['linkedin'])
        if contact.get('portfolio'):
            links.append(contact['portfolio'])

        if links:
            links_para = doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_run = links_para.add_run(' | '.join(links))
            links_run.font.size = Pt(10)
            links_run.font.name = self.font_name

        # Add spacing
        doc.add_paragraph()

    def _add_section(self, doc: Document, title: str) -> None:
        """Add section heading"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run(title)
        heading_run.font.size = self.font_size_heading
        heading_run.font.name = self.font_name
        heading_run.bold = True

        # Add underline manually (as bottom border)
        heading.paragraph_format.space_after = Pt(6)

    def _add_paragraph(self, doc: Document, text: str) -> None:
        """Add regular paragraph"""
        para = doc.add_paragraph(text)
        para_run = para.runs[0]
        para_run.font.size = self.font_size_body
        para_run.font.name = self.font_name
        para.paragraph_format.space_after = Pt(10)

    def _add_experience(self, doc: Document, exp: Dict) -> None:
        """Add work experience item"""
        # Job title (bold)
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(exp.get('title', ''))
        title_run.font.size = self.font_size_subheading
        title_run.font.name = self.font_name
        title_run.bold = True
        title_para.paragraph_format.space_after = Pt(2)

        # Company and dates (italic)
        company_parts = []
        if exp.get('company'):
            company_parts.append(exp['company'])
        if exp.get('location'):
            company_parts.append(exp['location'])

        company_line = ' | '.join(company_parts)

        if exp.get('start_date'):
            date_str = f"{exp['start_date']} - {exp.get('end_date', 'Present')}"
            company_line += f" | {date_str}"

        company_para = doc.add_paragraph()
        company_run = company_para.add_run(company_line)
        company_run.font.size = self.font_size_body
        company_run.font.name = self.font_name
        company_run.italic = True
        company_para.paragraph_format.space_after = Pt(4)

        # Bullets
        for bullet in exp.get('bullets', []):
            bullet_text = bullet.lstrip('• ').strip()
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = self.font_size_body
                run.font.name = self.font_name
            bullet_para.paragraph_format.space_after = Pt(2)

        # Spacing after experience
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_skills_category(self, doc: Document, category: str, skills: List[str]) -> None:
        """Add skills category"""
        skills_para = doc.add_paragraph()

        # Category name (bold)
        category_run = skills_para.add_run(f"{category}: ")
        category_run.font.size = self.font_size_body
        category_run.font.name = self.font_name
        category_run.bold = True

        # Skills (regular)
        skills_text = ", ".join(skills)
        skills_run = skills_para.add_run(skills_text)
        skills_run.font.size = self.font_size_body
        skills_run.font.name = self.font_name

        skills_para.paragraph_format.space_after = Pt(6)

    def _add_education(self, doc: Document, edu: Dict) -> None:
        """Add education item"""
        # Degree (bold)
        degree_text = f"{edu.get('degree', '')} in {edu.get('field', '')}"
        degree_para = doc.add_paragraph()
        degree_run = degree_para.add_run(degree_text)
        degree_run.font.size = self.font_size_subheading
        degree_run.font.name = self.font_name
        degree_run.bold = True
        degree_para.paragraph_format.space_after = Pt(2)

        # Institution and date
        inst_parts = []
        if edu.get('institution'):
            inst_parts.append(edu['institution'])
        if edu.get('graduation_date'):
            inst_parts.append(f"Graduated: {edu['graduation_date']}")

        if inst_parts:
            inst_para = doc.add_paragraph(' | '.join(inst_parts))
            for run in inst_para.runs:
                run.font.size = self.font_size_body
                run.font.name = self.font_name
            inst_para.paragraph_format.space_after = Pt(2)

        # GPA if present
        if edu.get('gpa'):
            gpa_para = doc.add_paragraph(f"GPA: {edu['gpa']}")
            for run in gpa_para.runs:
                run.font.size = self.font_size_body
                run.font.name = self.font_name
            gpa_para.paragraph_format.space_after = Pt(8)

    def _add_certification(self, doc: Document, cert: Dict) -> None:
        """Add certification item"""
        cert_parts = [cert.get('name', '')]

        if cert.get('issuer'):
            cert_parts.append(cert['issuer'])
        if cert.get('date'):
            cert_parts.append(cert['date'])

        cert_text = ' | '.join(cert_parts)

        cert_para = doc.add_paragraph(cert_text, style='List Bullet')
        for run in cert_para.runs:
            run.font.size = self.font_size_body
            run.font.name = self.font_name
        cert_para.paragraph_format.space_after = Pt(2)

    def get_docx_stream(self, resume_structure: Dict[str, Any]) -> io.BytesIO:
        """
        Get DOCX as BytesIO stream for FastAPI response

        Args:
            resume_structure: Resume data structure

        Returns:
            BytesIO stream containing DOCX
        """
        docx_bytes = self.generate_docx(resume_structure)
        docx_stream = io.BytesIO(docx_bytes)
        docx_stream.seek(0)

        return docx_stream
