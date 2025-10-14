"""OCR Service - Resume text extraction using Gemini"""

import google.generativeai as genai
import base64
from pathlib import Path
import os
from dotenv import load_dotenv

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

class OCRService:
    def __init__(self):
        self.model = genai.GenerativeModel('gemini-2.0-flash-exp')

    async def extract_resume_text(self, file_path: str) -> dict:
        """Extract all text and data from a resume (PDF/image/DOCX/DOC/TXT)"""

        file_ext = Path(file_path).suffix.lower()

        # Handle text-based formats directly
        if file_ext == '.txt':
            return await self._extract_from_txt(file_path)
        elif file_ext == '.docx':
            return await self._extract_from_docx(file_path)
        elif file_ext == '.doc':
            return await self._extract_from_doc(file_path)

        # Handle PDF and images with Gemini OCR
        return await self._extract_with_gemini(file_path)

    async def _extract_with_gemini(self, file_path: str) -> dict:
        """Extract using Gemini OCR (for PDF/images)"""

        # Read file and convert to base64
        with open(file_path, 'rb') as f:
            file_data = f.read()

        base64_data = base64.b64encode(file_data).decode('utf-8')

        # Determine MIME type
        file_ext = Path(file_path).suffix.lower()
        mime_type_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.pdf': 'application/pdf'
        }
        mime_type = mime_type_map.get(file_ext, 'application/octet-stream')

        prompt = """Extract ALL information from this resume with 100% accuracy.

Pay special attention to:
- Full name and contact information
- Professional summary/objective
- Work experience (company, title, dates, responsibilities, accomplishments)
- Education (school, degree, dates, GPA if present)
- Skills (technical, soft skills, languages)
- Certifications and awards
- Projects and publications
- Any quantifiable metrics (percentages, dollar amounts, numbers)

Return ONLY a JSON object with this structure:
{
  "personal_info": {
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "website": ""
  },
  "summary": "",
  "experience": [
    {
      "company": "",
      "title": "",
      "start_date": "",
      "end_date": "",
      "location": "",
      "responsibilities": [],
      "accomplishments": []
    }
  ],
  "education": [
    {
      "school": "",
      "degree": "",
      "field": "",
      "start_date": "",
      "end_date": "",
      "gpa": ""
    }
  ],
  "skills": {
    "technical": [],
    "languages": [],
    "soft_skills": []
  },
  "certifications": [],
  "awards": [],
  "projects": []
}

Extract EVERYTHING you see. Do not summarize or interpret - extract exactly as written."""

        # Generate content
        result = self.model.generate_content([
            {
                "inline_data": {
                    "mime_type": mime_type,
                    "data": base64_data
                }
            },
            prompt
        ])

        response_text = result.text

        # Clean up JSON (remove markdown code blocks if present)
        if response_text.startswith('```json'):
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif response_text.startswith('```'):
            response_text = response_text.split('```')[1].split('```')[0].strip()

        import json
        return json.loads(response_text)

    async def _extract_from_txt(self, file_path: str) -> dict:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            raw_text = f.read()

        # Use Claude to structure the text into resume format
        return await self._structure_text_with_claude(raw_text)

    async def _extract_from_docx(self, file_path: str) -> dict:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)

            # Extract all text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            raw_text = '\n'.join(paragraphs)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            raw_text += '\n' + cell.text

            # Use Claude to structure the text
            return await self._structure_text_with_claude(raw_text)

        except ImportError:
            raise Exception("python-docx not installed. Install with: pip install python-docx")

    async def _extract_from_doc(self, file_path: str) -> dict:
        """Extract text from DOC file (older Word format)"""
        try:
            import subprocess
            import tempfile

            # Try using antiword (must be installed on system)
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=10
            )

            if result.returncode == 0:
                raw_text = result.stdout
                return await self._structure_text_with_claude(raw_text)
            else:
                raise Exception("Could not extract text from .doc file. antiword failed.")

        except FileNotFoundError:
            # antiword not installed, try alternative method
            raise Exception(
                "antiword not installed. For .doc support, install antiword: "
                "brew install antiword (macOS) or apt-get install antiword (Linux)"
            )
        except Exception as e:
            raise Exception(f"Failed to extract .doc file: {str(e)}")

    async def _structure_text_with_claude(self, raw_text: str) -> dict:
        """Use Claude to structure raw text into resume format"""
        import anthropic

        claude = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY") or os.getenv("CLAUDE_API_KEY"))

        prompt = f"""Extract ALL information from this resume text and structure it properly.

RAW RESUME TEXT:
{raw_text}

Return ONLY a JSON object with this structure:
{{
  "personal_info": {{
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "website": ""
  }},
  "summary": "",
  "experience": [
    {{
      "company": "",
      "title": "",
      "start_date": "",
      "end_date": "",
      "location": "",
      "responsibilities": [],
      "accomplishments": []
    }}
  ],
  "education": [
    {{
      "school": "",
      "degree": "",
      "field": "",
      "start_date": "",
      "end_date": "",
      "gpa": ""
    }}
  ],
  "skills": {{
    "technical": [],
    "languages": [],
    "soft_skills": []
  }},
  "certifications": [],
  "awards": [],
  "projects": []
}}

Extract EVERYTHING you see. Do not summarize - extract exactly as written."""

        message = claude.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = message.content[0].text.strip()

        # Clean up JSON
        if response_text.startswith('```json'):
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif response_text.startswith('```'):
            response_text = response_text.split('```')[1].split('```')[0].strip()

        import json
        return json.loads(response_text)

# Create singleton instance
ocr_service = OCRService()
